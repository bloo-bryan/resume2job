import io

import docx


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)

    seen_cells: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text not in seen_cells:
                    seen_cells.add(cell.text)
                    parts.append(cell.text)

    return "\n".join(parts)
